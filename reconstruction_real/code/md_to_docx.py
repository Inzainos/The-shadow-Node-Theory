"""
Minimal Markdown -> Word (.docx) renderer for the SNT framework documents.
Handles: # / ## / ### headings, | pipe | tables, '- ' bullets, '> ' quotes,
4-space indented code blocks, and inline **bold** / *italic* / `code`.
Not a full Markdown engine — tuned for marco_teorico_v30(.md / _EN.md).

Usage: python md_to_docx.py input.md output.docx
"""
import re
import sys
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

INLINE = re.compile(r"(\*\*.+?\*\*|`[^`]+`|\*(?!\*).+?\*)")


def add_inline(p, text):
    """Add text to paragraph p, rendering **bold**, *italic*, `code`."""
    for tok in INLINE.split(text):
        if not tok:
            continue
        if tok.startswith("**") and tok.endswith("**"):
            p.add_run(tok[2:-2]).bold = True
        elif tok.startswith("`") and tok.endswith("`"):
            r = p.add_run(tok[1:-1]); r.font.name = "Consolas"; r.font.size = Pt(9)
        elif tok.startswith("*") and tok.endswith("*") and len(tok) > 2:
            p.add_run(tok[1:-1]).italic = True
        else:
            p.add_run(tok)


def flush_table(doc, rows):
    # rows: list of list[str]; drop separator rows like |---|---|
    rows = [r for r in rows if not all(set(c.strip()) <= set("-: ") for c in r)]
    if not rows:
        return
    ncol = max(len(r) for r in rows)
    t = doc.add_table(rows=0, cols=ncol)
    t.style = "Light Grid Accent 1"
    for ri, r in enumerate(rows):
        cells = t.add_row().cells
        for ci in range(ncol):
            txt = r[ci].strip() if ci < len(r) else ""
            cell_p = cells[ci].paragraphs[0]
            add_inline(cell_p, txt)
            if ri == 0:
                for run in cell_p.runs:
                    run.bold = True


def main(src, dst):
    md = open(src, encoding="utf-8").read().split("\n")
    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(10.5)

    tbuf, cbuf = [], []

    def flush_code():
        if cbuf:
            p = doc.add_paragraph()
            r = p.add_run("\n".join(cbuf))
            r.font.name = "Consolas"; r.font.size = Pt(9)
            cbuf.clear()

    for line in md:
        # table rows
        if line.strip().startswith("|") and line.strip().endswith("|"):
            flush_code()
            tbuf.append([c for c in line.strip().strip("|").split("|")])
            continue
        elif tbuf:
            flush_table(doc, tbuf); tbuf = []

        if line.startswith("    ") and line.strip():      # indented code
            cbuf.append(line[4:]); continue
        flush_code()

        s = line.rstrip()
        if not s.strip():
            continue
        if s.startswith("# "):
            doc.add_heading(s[2:], level=1)
        elif s.startswith("## "):
            doc.add_heading(s[3:], level=2)
        elif s.startswith("### "):
            doc.add_heading(s[4:], level=3)
        elif s.strip() in ("---", "***"):
            doc.add_paragraph()
        elif s.lstrip().startswith(("- ", "* ")):
            p = doc.add_paragraph(style="List Bullet")
            add_inline(p, s.lstrip()[2:])
        elif s.startswith(">"):
            p = doc.add_paragraph(); p.paragraph_format.left_indent = Pt(18)
            add_inline(p, s.lstrip("> ").rstrip())
            for r in p.runs:
                r.italic = True; r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        else:
            p = doc.add_paragraph()
            add_inline(p, s)

    if tbuf:
        flush_table(doc, tbuf)
    flush_code()
    doc.save(dst)
    print("DOCX written:", dst)


if __name__ == "__main__":
    main(sys.argv[1], sys.argv[2])
